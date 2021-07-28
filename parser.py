import docx
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed  # Error
import re


def pdfParse(file_name, save_path):
    pdf = open(f'instructions/{file_name}', 'r', encoding='utf-8')
    parser = PDFParser(pdf)
    doc = PDFDocument()
    parser.set_document(doc)
    doc.set_parser(parser)
    doc.initialize()
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        rsrcmagr = PDFResourceManager()
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmagr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmagr, device)
        for page in doc.get_pages():
            interpreter.process_page(page)
            layout = device.get_result()
            for x in layout:
                try:
                    if isinstance(x, LTTextBoxHorizontal):
                        with open(f'{save_path}', 'a', encoding='utf-8') as f:
                            result = x.get_text()
                            # print(result)
                            f.write(result + "\n")
                except:
                    print("Failed")


def htmlParse(html, save_path):
    pass


def docParse(file_name, save_path):  #形参：文件名，保存路径（包含文件名）
    f = open(f'{save_path}', 'a', encoding='utf-8')
    doc = docx.Document(f'instructions/{file_name}')
    # 两种组织形式 段落和表格，无法同时读取
    for para in doc.paragraphs:  # 读取基本文字
        f.write(para.text + '\n')
    f.write('----' + '\n')  # 区分文字段落与表格
    tables = doc.tables  # 获取文件中的表格集
    for table in tables[:]:
        result = []
        for i, row in enumerate(table.rows[:]):  # 读每行
            row_content = []
            for cell in row.cells[:]:  # 读一行中的所有单元格
                c = cell.text
                row_content.append(c)
            result.append(row_content)  # 将每一行组织起来
        # result 写入表格 还原顺序在解析中尝试
        f.write(str(result) + '\n')
    f.close()


class Parser:  # 解析当前目录下已经转换为txt的说明书
    def __init__(self, n, dosageForm, domestic=True):  # n:16进制顺序码 domestic是否为国产
        self.id = ''
        self.n = n
        self.dosageForm = dosageForm
        if domestic:
            self.domestic = 0
        else:
            self.domestic = 1

        f = open('demo.txt', 'r', encoding='utf-8')
        self.lines = f.readlines()  # 包含所有行的列表

        self.delimit = []  # 划分每一个部分，汇总到列表中
        delimiter = '^【.+】'  # 各个信息的界定符 （部分信息标题与内容在一行上展示）
        start = False
        para = ""  # 解决在赋值前引用的警告
        for line in self.lines:
            if re.match(delimiter, line):
                if start:
                    self.delimit.append(para)
                else:
                    start = True
                para = ""  # 清空字符串
            if start:
                para += line
        self.delimit.append(para.split('----')[0].strip()) # 并不写入表格数据，表格数据后期通过表头确定（取舍）

    def get_basic_info(self):
        basic_info = {}

        def find(item, field):
            for s in self.delimit:
                if re.match(f'^【{item}】\\n.+\\n$', s) or re.match(f'^【{item}】.+\\n$', s):
                    basic_info[field] = s.strip().lstrip(f'【{item}】\n')
                    break
                else:
                    basic_info[field] = 'NULL'

        find('批准文号', 'approvalNum')

        self.id = str(self.domestic)+basic_info['approvalNum'][4]+str(self.n)
        basic_info['id'] = self.id
        basic_info['dosageForm'] = self.dosageForm

        find('性状', 'character')
        basic_info['character'] = basic_info['character'].lstrip('本品为')
        find('规格', 'specification')
        find('执行标准', 'standard')
        find('有效期', 'validityTerm')
        basic_info['validityTerm'] = int(basic_info['validityTerm'].split('个月')[0])


if __name__ == "__main__":
    docParse('cde_domestic', )
    p = Parser(0, '片剂')
    p.get_basic_info()






